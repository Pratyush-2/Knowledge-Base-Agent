import os
import fitz  # PyMuPDF
import docx
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

def clean_text(text: str) -> str:
    """
    Cleans up common messy characters from text extraction.
    - Replaces special bullet points and characters with standard ones.
    - Removes excessive newlines and whitespace.
    """
    replacements = {
        '·': '-',
        '•': '-',
        '✔': '✓',
        '’': "'",
        '“': '"',
        '”': '"',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    text = '\n'.join(line.strip() for line in text.splitlines() if line.strip())
    
    return text

def _chunk_text(text: str, file_name: str, page_number: int = 1) -> List[Document]:
    """Splits text into chunks and creates Document objects."""
    cleaned_text = clean_text(text)
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=700,
        chunk_overlap=100,
        length_function=len,
        is_separator_regex=False,
    )
    
    text_chunks = text_splitter.split_text(cleaned_text)
    
    documents = []
    for i, chunk in enumerate(text_chunks):
        documents.append(Document(
            page_content=chunk,
            metadata={
                'source': file_name,
                'page': page_number,
                'chunk_number': i + 1
            }
        ))
        
    return documents

def process_uploaded_file(uploaded_file) -> List[Document]:
    """
    Detects the file type of an uploaded file and processes it into a list of
    chunked Document objects.
    """
    file_name = uploaded_file.name
    file_extension = os.path.splitext(file_name)[1].lower()
    
    all_chunks = []
    
    if file_extension == ".pdf":
        with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
            for i, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    chunks = _chunk_text(text, file_name, page_number=i + 1)
                    all_chunks.extend(chunks)
    elif file_extension == ".txt":
        text = uploaded_file.read().decode("utf-8")
        all_chunks = _chunk_text(text, file_name)
    elif file_extension == ".docx":
        doc = docx.Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        all_chunks = _chunk_text(text, file_name)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")
        
    return all_chunks